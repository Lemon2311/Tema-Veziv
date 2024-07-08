import os
from docx import Document

def parse(responsePath, outputPath):
    with open(responsePath, 'r', encoding='utf-8') as file:
        content = file.read()

    content = content.replace('\\n', '\n')
    start_index = content.find("Response:") + len("Response:")
    end_index = content.find("<|end_of_text|>']")

    doc = Document()

    if start_index != -1 and end_index != -1:
        substring = content[start_index:end_index].strip()
        lines = substring.split('\n')
        for line in lines:
            doc.add_paragraph(line)
    else:    
        doc.add_paragraph("The required text markers were not found in the file.")

    doc.save(outputPath)

def process_all_files(input_dir, output_dir):
    for filename in os.listdir(input_dir):
        if filename.endswith('.txt'):
            input_path = os.path.join(input_dir, filename)
            output_path = os.path.join(output_dir, filename.replace('.txt', '.docx'))
            parse(input_path, output_path)

input_dir = 'FineTunedModelOutputExamplesRaw'
output_dir = 'FineTunedModelOutputExamplesDocx'

if not os.path.exists(output_dir):
    os.makedirs(output_dir)

process_all_files(input_dir, output_dir)