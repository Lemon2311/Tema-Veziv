from docx import Document

document = Document("1.docx")

for paragraph in document.paragraphs:
    print(paragraph.text)   

# for paragraph in document.paragraphs:
#     # Paragraph style
#     style_name = paragraph.style.name
#     # Indentation (left indentation in Twips, convert to points: 1 point = 20 Twips)
#     indent = paragraph.paragraph_format.left_indent
#     indent_points = indent.pt if indent else 0  # Convert Twips to points if indent exists
#     # List style
#     is_list = paragraph.style.name.startswith('List')  # Simple check for list style
    
#     # Collecting text and formatting details
#     text_details = []
#     for run in paragraph.runs:
#         text = run.text
#         bold = "Bold" if run.bold else ""
#         text_details.append(f"{text} {bold}".strip())

#     text_details_str = ", ".join(text_details)
#     print(f"Style: {style_name}, Indent: {indent_points}pt, List: {is_list}, Text: {text_details_str}")

